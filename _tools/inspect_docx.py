"""Inspect the resume .docx so we can match styles when editing."""
import sys
from docx import Document

import sys
path = sys.argv[1] if len(sys.argv) > 1 else r"C:\Users\hamze\Desktop\Hamzeh_Emreish_Resume_Portfolio (1).docx"
doc = Document(path)

print(f"=== Document has {len(doc.paragraphs)} paragraphs ===\n")
for i, p in enumerate(doc.paragraphs):
    style = p.style.name if p.style else "(no style)"
    text = p.text.strip()
    # Truncate long lines for readability
    preview = text[:130] + ("..." if len(text) > 130 else "")
    print(f"[{i:3d}] style={style!r:30s} | {preview}")
