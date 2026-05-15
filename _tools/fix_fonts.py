"""
Fix font drift in the resume bullets I edited.

When edit_resume.py + bold_pass.py replaced or inserted runs, the new
runs were created without explicit <w:rPr><w:rFonts>...</w:rFonts>...,
so Word fell back to the paragraph-style default font — which doesn't
match the rest of the resume.

Fix: locate an untouched reference paragraph (a Technical Skills bullet),
extract its first run's font properties (rFonts / sz / szCs / color),
then apply those properties to every run in the paragraphs I edited.
Bold/italic on individual runs is preserved.
"""
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = r"C:\Users\hamze\Desktop\Hamzeh_Emreish_Resume_Portfolio_NEW.docx"

# Paragraph text starts I touched in edit_resume.py / bold_pass.py.
# Covers both the original named bullets and the anonymized rewrite.
TOUCHED_STARTS = [
    # BB heading (Normal style — also reset)
    "Bug Bounty Hunter HackerOne",
    # BB bullet bodies (List Paragraph)
    "Discovered a Critical (CVSS 9.1)",
    "Identified High-severity email injection",
    "Discovered a High-severity (CVSS 8.6)",
    "Found admin user enumeration",
    "Reported OAuth token exposure",
    "Discovered hardcoded production API credentials",
    "Identified unauthenticated mail relay",
    # Equinox subline
    "Promoted from Front Desk Associate; previously at Mamaroneck.",
    # Equinox bullets
    "Contribute to strategy meetings",
    "Rebuilt front desk schedule",
    "Run point on in-house tech",
    "Daily club walkthroughs",
    "Coach the front desk team",
]

# Reference paragraphs that were NOT touched and carry the canonical font.
REFERENCE_STARTS = [
    "Security Tools & Testing:",
    "Programming & Scripting:",
    "Developed and supported the full-stack",
    "Hired and trained a team",
]

# Font-related rPr child tags we want to mirror from the reference.
FONT_TAGS = ("w:rFonts", "w:sz", "w:szCs", "w:color", "w:lang")


def find_reference_rPr(doc):
    for p in doc.paragraphs:
        text = p.text.strip()
        if any(text.startswith(s) for s in REFERENCE_STARTS):
            for run in p.runs:
                rPr = run._element.find(qn("w:rPr"))
                if rPr is not None and rPr.find(qn("w:rFonts")) is not None:
                    return rPr
    return None


def ensure_rPr(run):
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        run._element.insert(0, rPr)
    return rPr


def apply_font(rPr_target, rPr_ref):
    for tag in FONT_TAGS:
        # Replace any existing of this tag with the reference's copy.
        for existing in list(rPr_target.findall(qn(tag))):
            rPr_target.remove(existing)
        ref_elem = rPr_ref.find(qn(tag))
        if ref_elem is not None:
            rPr_target.append(deepcopy(ref_elem))


def is_touched(para):
    text = para.text.strip()
    return any(text.startswith(s) for s in TOUCHED_STARTS)


def main():
    doc = Document(PATH)

    ref_rPr = find_reference_rPr(doc)
    if ref_rPr is None:
        print("ERROR: Could not find a reference paragraph with explicit fonts.")
        return

    print("Reference run-properties found:")
    for child in ref_rPr:
        tag_local = child.tag.split("}")[-1]
        print(f"  <w:{tag_local}> {dict(child.attrib)}")

    fixed_paras = 0
    fixed_runs = 0
    for p in doc.paragraphs:
        if not is_touched(p):
            continue
        preview = p.text.strip()[:65]
        if not p.runs:
            print(f"  SKIP (no runs): {preview!r}")
            continue
        for run in p.runs:
            rPr = ensure_rPr(run)
            apply_font(rPr, ref_rPr)
            fixed_runs += 1
        fixed_paras += 1
        print(f"  Fixed: {preview!r}")

    doc.save(PATH)
    print(f"\nApplied reference font to {fixed_runs} runs across {fixed_paras} paragraphs.")
    print(f"Saved: {PATH}")


if __name__ == "__main__":
    main()
