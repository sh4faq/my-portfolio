"""
Apply inline bold formatting to severity/status tags in the Bug Bounty bullets
of the NEW resume docx, matching the original formatting style.
"""
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = r"C:\Users\hamze\Desktop\Hamzeh_Emreish_Resume_Portfolio_NEW.docx"

# (anchor-substring-in-bullet, substring-to-bold)
# Anchor uniquely identifies which bullet; target is what to bold inside it.
BOLD_TARGETS = [
    ("enterprise observability platform's SQL expression",  "Critical (CVSS 9.1)"),
    ("enterprise LMS platform's Frappe-based contact form", "High-severity"),
    ("major cloud provider's GenAI Knowledge",              "High-severity (CVSS 8.6)"),
    ("Cognito ForgotPassword API",                          "(High, CVSS 7.5)"),
    ("major consumer fitness platform",                     "P5/Informational"),
    ("major DNS / internet infrastructure provider",        "(CWE-798, CWE-522)"),
    ("unauthenticated mail relay",                          "(Medium severity, triaged valid)"),
]


def add_run(elem, text, bold=False):
    r = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    elem.append(r)


def bold_substring(para, target):
    """Locate target in para.text and rebuild the paragraph as
    [plain pre-run, bold target run, plain post-run]. Returns True on success."""
    full = para.text
    idx = full.find(target)
    if idx == -1:
        return False
    pre = full[:idx]
    post = full[idx + len(target):]
    for r in list(para._element.findall(qn("w:r"))):
        para._element.remove(r)
    if pre:
        add_run(para._element, pre, bold=False)
    add_run(para._element, target, bold=True)
    if post:
        add_run(para._element, post, bold=False)
    return True


def main():
    doc = Document(PATH)
    applied = 0
    for para in doc.paragraphs:
        if para.style.name != "List Paragraph":
            continue
        text = para.text
        for anchor, target in BOLD_TARGETS:
            if anchor in text:
                ok = bold_substring(para, target)
                status = "OK" if ok else "NOT FOUND"
                print(f"  [{status}] anchor={anchor[:35]!r:38s} -> bold {target!r}")
                if ok:
                    applied += 1
                break
    doc.save(PATH)
    print(f"\nApplied {applied}/{len(BOLD_TARGETS)} bold targets. Saved to: {PATH}")


if __name__ == "__main__":
    main()
