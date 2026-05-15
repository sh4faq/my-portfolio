"""
Insert the Anthropic Cyber Verification Program (CVP) credential into the
resume's Certifications section as the FIRST bullet, cloning an existing
certification paragraph's XML so the bullet marker and font carry over
exactly.
"""
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

PATH = r"C:\Users\hamze\Desktop\Hamzeh_Emreish_Resume_Portfolio_NEW.docx"

CVP_TEXT = (
    "Anthropic Cyber Verification Program (CVP) — Approved May 14, 2026 "
    "for dual-use security research (vulnerability exploitation, offensive security tooling)"
)


def main():
    doc = Document(PATH)

    template = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("CompTIA Security+"):
            template = p
            break

    if template is None:
        print("ERROR: Could not find the CompTIA Security+ paragraph to use as template.")
        return

    if any(p.text.strip().startswith("Anthropic Cyber Verification") for p in doc.paragraphs):
        print("CVP already present — skipping insert.")
        return

    # Snapshot the run properties (font / size / bold-state) from the first run.
    template_rPr = None
    if template.runs:
        existing = template.runs[0]._element.find(qn("w:rPr"))
        if existing is not None:
            template_rPr = existing

    # Clone the entire paragraph XML so paragraph-level properties (bullet
    # numbering reference, indent, spacing) carry over to the new entry.
    new_elem = deepcopy(template._element)
    for r in list(new_elem.findall(qn("w:r"))):
        new_elem.remove(r)

    r = OxmlElement("w:r")
    if template_rPr is not None:
        r.append(deepcopy(template_rPr))
    t = OxmlElement("w:t")
    t.text = CVP_TEXT
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_elem.append(r)

    # Insert above the first existing cert so the CVP is the top bullet.
    template._element.addprevious(new_elem)

    doc.save(PATH)
    print(f"Inserted CVP credential as the first item in Certifications.")
    print(f"Saved: {PATH}")


if __name__ == "__main__":
    main()
