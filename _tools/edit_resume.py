"""
Edit Hamzeh's resume .docx:
  - Update Bug Bounty Hunter section (3 bullets -> 7) and add Intigriti to platforms.
  - Update Equinox section (subline + 4 bullets -> 5 bullets).

Strategy:
  - Modify heading/subline text by editing the existing run that contains the
    text-to-replace, preserving bold/italic formatting at the run level.
  - For bullet bodies, replace the entire paragraph's runs with a single plain
    run carrying the new text. Existing bullet bodies are plain (severity-tag
    bold is light cosmetic touch user can re-apply post-edit in Word).
  - For inserts, deepcopy the LAST existing bullet's XML element (so the new
    paragraphs inherit the same paragraph style, numbering id, and indent),
    clear its runs, and add a single new run with the new text.

Saves to a NEW file so we can sanity-check in Word before overwriting.
"""
from copy import deepcopy
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SOURCE = r"C:\Users\hamze\Desktop\Hamzeh_Emreish_Resume_Portfolio (1).docx"
OUTPUT = r"C:\Users\hamze\Desktop\Hamzeh_Emreish_Resume_Portfolio_NEW.docx"

# ---------------------------------------------------------------------------
# NEW CONTENT
# ---------------------------------------------------------------------------

BB_BULLETS = [
    "Discovered a Critical (CVSS 9.1) parser allowlist bypass in an enterprise observability platform's SQL expression engine (Intigriti); the fix for CVE-2026-27876 left 4 of 5 SetOp AST fields skipped by vitess walkSubtree unguarded, enabling Local File Inclusion via load_file() inside an ORDER BY clause. Same primitive class as CVE-2024-9264. Triaged as duplicate.",
    "Identified High-severity email injection on an enterprise LMS platform's Frappe-based contact form endpoint (HackerOne); the sender parameter is used as the email recipient rather than the sender address, enabling attacker-controlled outbound mail from the platform's legitimate domain to arbitrary recipients with full SPF/DKIM/DMARC alignment and no rate limit. Working end-to-end PoC with chained phishing payload; submitted, awaiting triage.",
    "Discovered a High-severity (CVSS 8.6) SSRF in a major cloud provider's GenAI Knowledge Base web crawler (Intigriti); HTTP 302 redirect bypass of hostname deny-list and IP resolution check reaches 169.254.169.254 cloud metadata service and internal infrastructure. Submitted May 14, 2026; awaiting triage.",
    "Found admin user enumeration on an enterprise LMS platform via Cognito ForgotPassword API (High, CVSS 7.5) on HackerOne; unauthenticated GraphQL query leaks the admin client_id, and Cognito's ForgotPassword endpoint returns distinct responses for valid versus invalid usernames (CodeDeliveryDetails vs UserNotFoundException), enabling admin email enumeration with first-letter leak. Patched shortly after testing; triaged as duplicate.",
    "Reported OAuth token exposure on a major consumer fitness platform's API via Bugcrowd; access tokens embedded in window.__STATE__ on the public OAuth page enabled authentication bypass with confirmed account modification (PUT to a user resource returned HTTP 200, first_name modified) and mass user enumeration across IDs 1 through 231M+ via email, name, and query searches without rate limiting. Program ruled the modified account a test user and classified P5/Informational.",
    "Discovered hardcoded production API credentials (CWE-798, CWE-522) in a publicly-accessible JavaScript config on a major DNS / internet infrastructure provider's domain-suggestion product (Bugcrowd), granting unauthenticated access to all seven backend API endpoints. Triaged as duplicate of an earlier report (April 2025) with restricted key scope.",
    "Identified unauthenticated mail relay vulnerability (Medium severity, triaged valid) enabling email spoofing with valid DKIM/SPF signatures; provided CVSS analysis and remediation guidance.",
]

EQ_SUBLINE_NEW = "Promoted from Front Desk Associate; previously at Mamaroneck. Flagship rotations at Hudson Yards and Printing House."

EQ_BULLETS = [
    "Contribute to strategy meetings with the General Manager, Regional Training Manager, and Facility Manager on member retention, personal training conversion, facility quality, and staff performance.",
    "Rebuilt front desk schedule and shift structure; closed a $1,500 labor budget overrun and brought it to zero within three months.",
    "Run point on in-house tech for the club (BitLocker recovery, Windows blue screens, POS configuration, audio system faults, member-facing hardware); diagnose first, escalate to the vendor only when a part has genuinely failed.",
    "Daily club walkthroughs surface cleanliness and equipment risk; tag failing gear Out of Service and power-down treadmills at the source. Caught a member accessing the club on someone else's account for nearly a month from a photo-ID mismatch the desk had missed; escalated to GM, who verified and took action.",
    "Coach the front desk team on luxury-club engagement standards (phones off the desk, iPad-only check-in, security verification, greeting every member like a guest). Trained and onboarded 10+ team members; cross-trained in Pro Shop. CPR/AED certified.",
]

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def replace_in_runs(para, old_substr, new_substr):
    """Find a run containing old_substr and replace it in-place. Preserves
    bold/italic formatting at the run level. Returns True if substitution
    happened, False otherwise."""
    for run in para.runs:
        if old_substr in run.text:
            run.text = run.text.replace(old_substr, new_substr)
            return True
    # Fallback: concatenated text doesn't span a single run. Try across all runs
    # by clearing all runs and rebuilding a single plain run with the substitution.
    full = para.text
    if old_substr in full:
        new_full = full.replace(old_substr, new_substr)
        clear_paragraph_runs(para)
        add_run(para._element, new_full)
        return True
    return False


def clear_paragraph_runs(para):
    """Remove all <w:r> children from the paragraph element."""
    for r in list(para._element.findall(qn("w:r"))):
        para._element.remove(r)


def add_run(elem, text):
    """Append a single <w:r><w:t>text</w:t></w:r> to elem (xml:space=preserve)."""
    r = OxmlElement("w:r")
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    elem.append(r)


def replace_para_text(para, new_text):
    """Replace all runs in para with a single plain run carrying new_text."""
    clear_paragraph_runs(para)
    add_run(para._element, new_text)


def clone_bullet_with_text(template_para, new_text):
    """Deepcopy template_para's XML, clear its runs, add new_text as a single run.
    Returns the new XML element ready to be inserted."""
    new_elem = deepcopy(template_para._element)
    for r in list(new_elem.findall(qn("w:r"))):
        new_elem.remove(r)
    add_run(new_elem, new_text)
    return new_elem


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    doc = Document(SOURCE)
    paras = doc.paragraphs

    bb_heading_p = None
    eq_heading_p = None
    eq_subline_p = None
    bb_bullet_paras = []
    eq_bullet_paras = []

    in_bb = False
    in_eq = False
    for p in paras:
        txt = p.text.strip()
        if txt.startswith("Bug Bounty Hunter"):
            bb_heading_p = p
            in_bb, in_eq = True, False
        elif txt.startswith("Equinox"):
            eq_heading_p = p
            in_bb, in_eq = False, True
        elif "Promoted from Front Desk Associate" in txt:
            eq_subline_p = p
        elif txt.startswith("Yonkers Car Wash"):
            in_bb, in_eq = False, False
        elif p.style.name == "List Paragraph":
            if in_bb:
                bb_bullet_paras.append(p)
            elif in_eq:
                eq_bullet_paras.append(p)

    assert bb_heading_p is not None, "Bug Bounty heading not found"
    assert eq_heading_p is not None, "Equinox heading not found"
    assert eq_subline_p is not None, "Equinox subline not found"
    assert len(bb_bullet_paras) == 3, f"Expected 3 BB bullets, got {len(bb_bullet_paras)}"
    assert len(eq_bullet_paras) == 4, f"Expected 4 Eq bullets, got {len(eq_bullet_paras)}"

    print("Anchors found:")
    print(f"  BB heading       : {bb_heading_p.text!r}")
    print(f"  BB bullets       : {len(bb_bullet_paras)}")
    print(f"  Equinox heading  : {eq_heading_p.text!r}")
    print(f"  Equinox subline  : {eq_subline_p.text!r}")
    print(f"  Equinox bullets  : {len(eq_bullet_paras)}")

    # -----------------------------------------------------------------------
    # BUG BOUNTY SECTION
    # -----------------------------------------------------------------------
    print("\nUpdating Bug Bounty section...")

    # Add " / Intigriti" to the platforms portion of the heading.
    # The original platforms substring is "HackerOne / Bugcrowd".
    ok = replace_in_runs(bb_heading_p, "HackerOne / Bugcrowd", "HackerOne / Bugcrowd / Intigriti")
    print(f"  Heading update ok? {ok}")

    # Replace existing 3 bullet bodies with the first 3 new ones.
    for existing, new_text in zip(bb_bullet_paras, BB_BULLETS[:3]):
        replace_para_text(existing, new_text)
    print(f"  Replaced 3 existing BB bullets")

    # Insert remaining 4 bullets after the last existing one. Iterate in reverse
    # so each addnext lands at the correct position relative to last_existing.
    last_existing = bb_bullet_paras[-1]
    for new_text in reversed(BB_BULLETS[3:]):
        new_elem = clone_bullet_with_text(last_existing, new_text)
        last_existing._element.addnext(new_elem)
    print(f"  Inserted {len(BB_BULLETS) - 3} new BB bullets")

    # -----------------------------------------------------------------------
    # EQUINOX SECTION
    # -----------------------------------------------------------------------
    print("\nUpdating Equinox section...")

    # Update the italic subline (whole-text replacement preserves italic on the run).
    ok = replace_in_runs(
        eq_subline_p,
        "Promoted from Front Desk Associate; previously at Mamaroneck location",
        EQ_SUBLINE_NEW,
    )
    if not ok:
        # Fallback: the original might be slightly different; do a whole-paragraph replace.
        replace_para_text(eq_subline_p, EQ_SUBLINE_NEW)
        ok = True
    print(f"  Subline update ok? {ok}")

    # Replace existing 4 bullet bodies with first 4 new ones.
    for existing, new_text in zip(eq_bullet_paras, EQ_BULLETS[:4]):
        replace_para_text(existing, new_text)
    print(f"  Replaced 4 existing Eq bullets")

    # Insert 1 new bullet (5th) after the last existing one.
    last_existing = eq_bullet_paras[-1]
    new_elem = clone_bullet_with_text(last_existing, EQ_BULLETS[4])
    last_existing._element.addnext(new_elem)
    print(f"  Inserted 1 new Eq bullet")

    # Save
    doc.save(OUTPUT)
    print(f"\nSaved to: {OUTPUT}")


if __name__ == "__main__":
    main()
